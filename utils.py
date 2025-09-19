# -*- coding: utf-8 -*-
# utils.py — Parte 1/4 (Base + Ingesta KB + Router OpenAI) — versión mejorada

"""
Módulo utilitario para:
- Extracción de texto (PDF/DOCX/Imágenes + OCR selectivo con OpenAI Vision)
- Normalizaciones y helpers varios
- Pipeline de análisis con modelos OpenAI
- Ingesta de Base de Conocimiento (KB): subir/leer archivos, trocear y embebidos
- RAG liviano (en Partes 3 y 4) sin dependencia de numpy.

Notas:
- Se usa lazy-import para evitar ciclos (p.ej. con prompts.py o modelos SQLAlchemy).
- Wrapper de Responses API (reemplaza chat.completions) con fallback seguro.
- Este archivo se entrega en 4 partes para facilitar el pegado.
"""

from __future__ import annotations

import io
import os
import re
import json
import time
import base64
import shutil
import hashlib
import mimetypes
from datetime import datetime
from typing import List, Tuple, Dict, Optional, Iterator, Any
from tempfile import NamedTemporaryFile
from concurrent.futures import ThreadPoolExecutor, as_completed

import fitz  # PyMuPDF
from dotenv import load_dotenv
from openai import OpenAI
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.pdfgen import canvas
from reportlab.lib.colors import HexColor
from reportlab.lib.utils import ImageReader
from zoneinfo import ZoneInfo

# ========================= Carga de .env =========================
load_dotenv()

# ========================= OpenAI client =========================
# Importante: el timeout del cliente global no siempre se aplica a cada request.
# Creamos un alias "client_timed" con timeout efectivo por request.
OPENAI_TIMEOUT = float(os.getenv("OPENAI_TIMEOUT", "90"))
_API_KEY = os.getenv("OPENAI_API_KEY")  # Render -> Environment Variables
_API_BASE = os.getenv("OPENAI_API_BASE") or os.getenv("OPENAI_BASE_URL") or None
client = OpenAI(api_key=_API_KEY, base_url=_API_BASE) if _API_BASE else OpenAI(api_key=_API_KEY)
client_timed = client.with_options(timeout=OPENAI_TIMEOUT)

# ========================= Rutear Responses -> Chat (helper) =========================
def _to_messages(val):
    """
    Normaliza `messages`/`input` a lista de {role, content (str)} para Chat fallback.
    Acepta string, lista de dicts con content str o multimodal (lo reduce a texto).
    """
    if not val:
        return [{"role": "user", "content": " "}]

    if isinstance(val, str):
        return [{"role": "user", "content": val}]

    norm = []
    for m in val:
        role = m.get("role") or "user"
        content = m.get("content")
        if isinstance(content, list):
            parts = []
            for p in content:
                ptype = (p.get("type") or "").lower()
                if ptype in ("text", "input_text") and p.get("text"):
                    parts.append(p["text"])
                elif ptype in ("image_url", "input_image", "image"):
                    raw = p.get("image_url") or p.get("image") or {}
                    url = raw.get("url") if isinstance(raw, dict) else (raw if isinstance(raw, str) else "")
                    if url:
                        parts.append(f"[imagen: {url}]")
            norm.append({"role": role, "content": "\n".join([x for x in parts if x])})
        else:
            norm.append({"role": role, "content": str(content or "")})
    return norm


def _responses_or_chat(_client, **attempt):
    """
    Usa Responses si está permitido y disponible; si OPENAI_FORCE_CHAT=1 o falla,
    cae a Chat Completions mapeando el payload.
    - Entradas posibles: model, input|messages, temperature, max_output_tokens|max_tokens|max_completion_tokens, metadata, tools, tool_choice.
    - Devuelve el objeto de respuesta del SDK (no el texto).
    """
    force_chat = os.getenv("OPENAI_FORCE_CHAT", "0") == "1"

    model_responses = attempt.get("model") or os.getenv("OPENAI_RESPONSES_MODEL", "gpt-4.1-mini")
    model_chat = os.getenv("OPENAI_CHAT_MODEL", "gpt-4o-mini")

    msgs = attempt.get("input") or attempt.get("messages")
    messages_for_chat = _to_messages(msgs)

    temperature = attempt.get("temperature", None)

    tok = attempt.pop("max_output_tokens", None) or attempt.pop("max_completion_tokens", None) or attempt.pop("max_tokens", None)
    try:
        tok = int(tok) if tok is not None else None
    except Exception:
        tok = None

    passthrough_keys = ("tools", "tool_choice", "metadata")

    if not force_chat:
        # --------- Intento Responses ---------
        r_payload = {"model": model_responses, "input": msgs}
        if temperature is not None:
            r_payload["temperature"] = float(temperature)
        if tok is not None:
            r_payload["max_output_tokens"] = tok
        for k in passthrough_keys:
            if k in attempt:
                r_payload[k] = attempt[k]

        try:
            return _client.responses.create(**r_payload)
        except TypeError:
            # SDK que no acepta algún campo: retirarlo y reintentar
            r_payload.pop("max_output_tokens", None)
            r_payload.pop("metadata", None)
            try:
                return _client.responses.create(**r_payload)
            except Exception:
                pass
        except Exception:
            # cualquier otro error -> probar Chat
            pass

    # --------- Fallback Chat Completions ---------
    c_payload = {"model": model_chat, "messages": messages_for_chat}
    if temperature is not None:
        c_payload["temperature"] = float(temperature)
    if tok is not None:
        c_payload["max_tokens"] = tok
    for k in passthrough_keys:
        if k in attempt:
            c_payload[k] = attempt[k]

    return _client.chat.completions.create(**c_payload)

# ========================= Embeddings model ======================
# Para costo/latencia: por defecto usamos text-embedding-3-small (suficiente para RAG liviano).
EMBEDDINGS_MODEL = (
    os.getenv("OPENAI_EMBEDDINGS_MODEL", "text-embedding-3-small").strip()
    or "text-embedding-3-small"
)

# ========================= Base de Conocimiento (KB) =========================
# ATENCIÓN: estas funciones usan los modelos definidos en models.py.
# Se importan de forma diferida para no registrar las tablas dos veces.
def _kb_models():
    try:
        from models import (
            KBSource as _KBSource,
            KBFile as _KBFile,
            KBChunk as _KBChunk,
            KBPriority as _KBPriority,
        )
    except Exception:
        # fallback si se ejecuta como paquete
        from .models import (  # type: ignore
            KBSource as _KBSource,
            KBFile as _KBFile,
            KBChunk as _KBChunk,
            KBPriority as _KBPriority,
        )
    return _KBSource, _KBFile, _KBChunk, _KBPriority


def _kb_clean_text(s: str) -> str:
    return " ".join((s or "").replace("\r", " ").replace("\n", " ").split())


def _kb_chunk_text(text: str, max_chars: int = 1500, overlap: int = 200) -> Iterator[str]:
    """Corta texto en bloques superpuestos para embebidos."""
    if not text:
        return iter(())
    n = len(text)
    i = 0
    while i < n:
        j = min(i + max_chars, n)
        yield text[i:j]
        if j >= n:
            break
        i = max(0, j - overlap)


def _kb_sha256_path(path: str) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for b in iter(lambda: f.read(1024 * 1024), b""):
            h.update(b)
    return h.hexdigest()


def _kb_detect_content_type(filename: str) -> str:
    fn = filename.lower()
    if fn.endswith(".pdf"):
        return "application/pdf"
    if fn.endswith(".json"):
        return "application/json"
    if fn.endswith((".txt", ".md", ".log", ".csv")):
        return "text/plain"
    return mimetypes.guess_type(filename)[0] or "text/plain"


def _kb_extract_text_from_path(path: str) -> Tuple[str, Dict[str, Any]]:
    """Extracción mínima para ingesta offline (la extracción ‘seria’ está más abajo)."""
    ext = os.path.splitext(path)[1].lower()
    meta: Dict[str, Any] = {}
    if ext in [".txt", ".md", ".csv", ".log"]:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read(), meta
    if ext == ".json":
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            obj = json.load(f)
        meta["json_keys"] = list(obj.keys()) if isinstance(obj, dict) else None
        return json.dumps(obj, ensure_ascii=False), meta
    if ext == ".pdf":
        try:
            with fitz.open(path) as doc:
                texts: List[str] = []
                spans: List[Dict[str, Any]] = []
                for p in range(len(doc)):
                    t = (doc[p].get_text("text") or "").strip()
                    texts.append(t)
                    spans.append({"page": p + 1, "chars": len(t)})
                meta["pages"] = len(doc)
                meta["spans"] = spans
                return "\n".join(texts), meta
        except Exception:
            return "", {"unsupported": True, "ext": ext, "error": "pdf_read_error"}
    return "", {"unsupported": True, "ext": ext}


def _kb_embed(text: str, _client: Optional[OpenAI] = None) -> List[float]:
    """
    Embedding con timeout efectivo por request.
    Sugerencia de costo: usar EMBEDDINGS_MODEL = text-embedding-3-small salvo que necesites más recall.
    """
    cli = _client or client_timed
    resp = cli.embeddings.create(model=EMBEDDINGS_MODEL, input=text)
    return resp.data[0].embedding  # type: ignore


def kb_create_or_get_source(db, name: str, storage_path: str, scope: Dict[str, Any]):
    """
    Crea (si no existe) o actualiza una fuente/rubro en la KB.
    - db: sesión SQLAlchemy del sistema (models.py)
    - name: nombre legible de la fuente
    - storage_path: carpeta donde se guardarán los archivos subidos
    - scope: metadatos/alcance (dict)
    """
    KBSource, _, _, _ = _kb_models()
    src = db.query(KBSource).filter_by(name=name).first()
    if src:
        changed = False
        if storage_path and src.storage_path != storage_path:
            src.storage_path = storage_path
            changed = True
        if scope and (src.scope or {}) != scope:
            src.scope = scope
            changed = True
        if changed:
            db.add(src)
            db.commit()
            db.refresh(src)
        return src
    os.makedirs(storage_path or ".", exist_ok=True)
    src = KBSource(name=name, storage_path=storage_path, scope=scope or {})
    db.add(src)
    db.commit()
    db.refresh(src)
    return src


def kb_ingest_file(
    db,
    source,
    local_path: str,
    *,
    rubric: str,
    tags: Optional[List[str]] = None,
    _client: Optional[OpenAI] = None,
    chunk_chars: int = 1500,
    overlap: int = 200,
) -> int:
    """
    Ingesta un archivo en la KB:
    - Copia el archivo a storage_path.
    - Extrae texto básico (rápido).
    - Trocea y guarda chunks + embeddings (como JSON).
    Devuelve el ID del KBFile.
    """
    KBSource, KBFile, KBChunk, _ = _kb_models()
    if not os.path.isfile(local_path):
        raise FileNotFoundError(f"No existe el archivo: {local_path}")
    os.makedirs(source.storage_path or ".", exist_ok=True)

    # Evitar duplicados por hash
    sha = _kb_sha256_path(local_path)
    existing = db.query(KBFile).filter_by(hash_sha256=sha).first()
    if existing:
        return existing.id

    dest = os.path.join(source.storage_path or ".", os.path.basename(local_path))
    if os.path.abspath(dest) != os.path.abspath(local_path):
        shutil.copy2(local_path, dest)

    size = os.path.getsize(dest)
    ctype = _kb_detect_content_type(dest)
    text, meta_extra = _kb_extract_text_from_path(dest)

    kbfile = KBFile(
        source_id=source.id,
        filename=os.path.basename(dest),
        content_type=ctype,
        bytes=size,
        hash_sha256=sha,
        stored_path=dest,
        meta={"rubro": rubric, "tags": (tags or []), **(meta_extra or {})},
    )
    db.add(kbfile)
    db.flush()

    # Indexar chunks si hay texto
    if (text or "").strip():
        text_norm = _kb_clean_text(text)
        ord_idx = 0
        for chunk in _kb_chunk_text(text_norm, max_chars=chunk_chars, overlap=overlap):
            vec = _kb_embed(chunk, _client=_client)
            db.add(
                KBChunk(
                    file_id=kbfile.id,
                    ord=ord_idx,
                    text=chunk,
                    embedding=json.dumps(vec),  # guardado como JSON string
                    span={},
                    meta={},
                )
            )
            ord_idx += 1

    db.commit()
    return kbfile.id


def kb_upsert_priority(db, rubric: str, label: str, details: str = "", weight: float = 1.0) -> None:
    """
    Crea o actualiza una prioridad de búsqueda en la KB (por rubro/etiqueta).
    Se usa luego para aumentar score durante la recuperación (RAG).
    """
    _, _, _, KBPriority = _kb_models()
    from sqlalchemy import and_

    row = (
        db.query(KBPriority)
        .filter(and_(KBPriority.rubric == rubric, KBPriority.label == label))
        .first()
    )
    if not row:
        row = KBPriority(rubric=rubric, label=label, details=details or "", weight=weight)
        db.add(row)
    else:
        row.details = details or row.details
        row.weight = weight
    db.commit()
# -*- coding: utf-8 -*-
# utils.py — Parte 2/4 (Prompts + Extracción base + Wrapper Responses→texto)

# ========================= Prompts (lazy import) =========================
_prom = None

def _get_prom():
    """Carga prompts.py bajo demanda, evitando ciclos de import."""
    global _prom
    if _prom is None:
        import importlib
        try:
            _prom = importlib.import_module("prompts")
        except Exception:
            try:
                pkg = __package__
                if pkg:
                    _prom = importlib.import_module(f"{pkg}.prompts")
                else:
                    _prom = None
            except Exception:
                _prom = None
    return _prom

def _sinonimos_text() -> str:
    """
    Texto opcional de sinónimos/convenciones. Si no existe en prompts.py,
    devuelve cadena vacía.
    """
    mod = _get_prom()
    return getattr(mod, "SINONIMOS_CANONICOS", "") if mod else ""

def _craft_prompt_notas_text() -> str:
    """
    Prompt breve para generar 'notas parciales' en la pasada multi-chunk.
    Si no existe en prompts.py, se usa un fallback sobrio.
    """
    mod = _get_prom()
    return getattr(
        mod,
        "CRAFT_PROMPT_NOTAS",
        "Extrae bullets técnicos y concisos con citas literales; cero invenciones."
    )

def _default_reglas_citas(varios_anexos: bool) -> str:
    if varios_anexos:
        return (
            "Reglas de Citas:\n"
            "- Documento MULTI-ANEXO: al final de cada línea con dato, usar (Anexo X, p. N).\n"
            "- Deducir N utilizando la etiqueta [PÁGINA N] más cercana dentro del texto del ANEXO correspondiente.\n"
            "- Si no hay paginación: (Fuente: documento provisto)."
        )
    else:
        return (
            "Reglas de Citas:\n"
            "- Documento ÚNICO: al final de cada línea con dato, usar (p. N) a partir de la etiqueta [PÁGINA N] más cercana.\n"
            "- Si no hay paginación: (Fuente: documento provisto)."
        )

def prompt_andres(varios_anexos: bool) -> str:
    """
    Devuelve el prompt maestro del ANALIZADOR, compatible con:
    - PROMPT_ANALIZADOR (nuevo)
    - PROMPT_ANALISIS (alias opcional)
    - PROMPT_PARAMETRIZADO (legacy)
    Inserta reglas de citas dinámicas y la NO_RENGLONES_RULE si está disponible.
    """
    mod = _get_prom()
    if mod:
        # 1) Localizar cuerpo de prompt del analizador (nuevo -> legacy)
        raw_prompt = None
        for key in ("PROMPT_ANALIZADOR", "PROMPT_ANALISIS", "PROMPT_PARAMETRIZADO"):
            if hasattr(mod, key):
                raw_prompt = getattr(mod, key)
                break

        # 2) Reglas de citas (función externa si existe, si no: default)
        if hasattr(mod, "reglas_citas"):
            reglas = mod.reglas_citas(varios_anexos)  # type: ignore[attr-defined]
        else:
            reglas = _default_reglas_citas(varios_anexos)

        # 3) Regla NO_RENGLONES (si no está, usar una segura por defecto)
        no_renglones = getattr(
            mod,
            "NO_RENGLONES_RULE",
            ("Para el campo 'Número de renglón' en la Ficha, escribir exactamente: "
             "'Total de renglones: <cantidad>; ver Sección 9 para el detalle completo'. "
             "Nunca uses 'N' como placeholder ni inventes cantidades.")
        )

        if raw_prompt:
            try:
                return raw_prompt.format(
                    REGLAS_CITAS=reglas,
                    NO_RENGLONES_RULE=no_renglones
                )
            except Exception:
                # Si el prompt no tiene placeholders, concatenamos reglas al final.
                return (str(raw_prompt).rstrip() + "\n\n" + reglas + "\n\n" + no_renglones).strip()

    # Fallback ultra-minimal si faltara prompts.py
    return (
        "Elabora un informe técnico-jurídico estructurado con citas literales. "
        "No inventes. Cita como '(p. N)'. Si hay múltiples anexos, usa '(Anexo X, p. N)'."
    )

# ========================= Modelos / Heurísticas =========================
MODEL_ANALISIS   = os.getenv("OPENAI_MODEL_ANALISIS", "gpt-4o-mini")
VISION_MODEL     = os.getenv("OPENAI_MODEL_VISION", "gpt-4o-mini")
MODEL_NOTAS      = os.getenv("OPENAI_MODEL_NOTAS", MODEL_ANALISIS)
MODEL_SINTESIS   = os.getenv("OPENAI_MODEL_SINTESIS", MODEL_ANALISIS)
FAST_FORCE_MODEL = os.getenv("FAST_FORCE_MODEL", "").strip()  # opcional para fast
FALLBACK_MODEL_DEFAULT = os.getenv("OPENAI_MODEL_FALLBACK", "gpt-4o-mini")

MAX_SINGLE_PASS_CHARS       = int(os.getenv("MAX_SINGLE_PASS_CHARS", "120000"))
MAX_SINGLE_PASS_CHARS_MULTI = int(os.getenv("MAX_SINGLE_PASS_CHARS_MULTI", str(MAX_SINGLE_PASS_CHARS)))

CHUNK_SIZE_BASE = int(os.getenv("CHUNK_SIZE", "24000"))
TARGET_PARTS    = int(os.getenv("TARGET_PARTS", "2"))

# Tope "blando" de tokens de salida (el wrapper puede reducir dinámicamente)
MAX_COMPLETION_TOKENS_SALIDA = int(os.getenv("MAX_COMPLETION_TOKENS_SALIDA", "3500"))
TEMPERATURE_ANALISIS         = os.getenv("TEMPERATURE_ANALISIS", "").strip()
ANALISIS_MODO                = os.getenv("ANALISIS_MODO", "").lower().strip()  # "fast" opcional

# Granularidad / anti-copia ligera
RENGLON_DESC_MAX_WORDS = int(os.getenv("RENGLON_DESC_MAX_WORDS", "24"))
ART_SNIPPET_MAX_WORDS  = int(os.getenv("ART_SNIPPET_MAX_WORDS", "18"))

# Concurrencia
ANALISIS_CONCURRENCY = int(os.getenv("ANALISIS_CONCURRENCY", "3"))
NOTAS_MAX_TOKENS     = int(os.getenv("NOTAS_MAX_TOKENS", "1400"))

# OCR
VISION_MAX_PAGES     = int(os.getenv("VISION_MAX_PAGES", "8"))
VISION_DPI           = int(os.getenv("VISION_DPI", "150"))
OCR_TEXT_MIN_CHARS   = int(os.getenv("OCR_TEXT_MIN_CHARS", "120"))
OCR_CONCURRENCY      = int(os.getenv("OCR_CONCURRENCY", "4"))

# Control de paginado en texto nativo
PAGINAR_TEXTO_NATIVO = int(os.getenv("PAGINAR_TEXTO_NATIVO", "1"))

# Calidad/recall
MULTI_FORCE_TWO_STAGE_MIN_CHARS = int(os.getenv("MULTI_FORCE_TWO_STAGE_MIN_CHARS", "45000"))
ENABLE_REGEX_HINTS              = int(os.getenv("ENABLE_REGEX_HINTS", "1"))
HINTS_MAX_CHARS                 = int(os.getenv("HINTS_MAX_CHARS", "12000"))
HINTS_PER_FIELD                 = int(os.getenv("HINTS_PER_FIELD", "8"))
ENABLE_SECOND_PASS_COMPLETION   = int(os.getenv("ENABLE_SECOND_PASS_COMPLETION", "1"))

# Ampliaciones automáticas
EXPAND_SECTIONS_213_216        = int(os.getenv("EXPAND_SECTIONS_213_216", "0"))
MAX_RENGLONES_OUT              = int(os.getenv("MAX_RENGLONES_OUT", "12"))
MAX_ARTICULOS_OUT              = int(os.getenv("MAX_ARTICULOS_OUT", "12"))
FORCE_DETERMINISTIC_213_216    = int(os.getenv("FORCE_DETERMINISTIC_213_216", "0"))

# ====== Gobernanza de longitud / orden de salida ======
STRICT_OUT             = int(os.getenv("STRICT_OUT", "1"))  # 1 = aplicar recortes y orden forzado
MAX_TOTAL_CHARS_OUT    = int(os.getenv("MAX_TOTAL_CHARS_OUT", "16000"))
MAX_LINES_PER_SECTION  = int(os.getenv("MAX_LINES_PER_SECTION", "20"))
MAX_WORDS_PER_BULLET   = int(os.getenv("MAX_WORDS_PER_BULLET", "35"))
SECTION_CHAR_LIMIT     = int(os.getenv("SECTION_CHAR_LIMIT", "2200"))
MAX_WORDS_TOTAL_GUIDE  = int(os.getenv("MAX_WORDS_TOTAL_GUIDE", "1200"))  # guía para el prompt

# ========================= Timers PERF =========================
def _t() -> float:
    return time.perf_counter()

def _log_tiempo(etiqueta: str, t0: float) -> None:
    try:
        dt = time.perf_counter() - t0
        print(f"[PERF] {etiqueta}: {dt:0.2f}s")
    except Exception:
        pass

# ==================== OCR / Raster ====================
def _rasterizar_pagina(page, dpi: int = VISION_DPI) -> bytes:
    mat = fitz.Matrix(dpi / 72, dpi / 72)
    pix = page.get_pixmap(matrix=mat, alpha=False)
    return pix.tobytes("png")

# ==================== Wrapper Responses API → texto ====================
def _resp_to_text(resp) -> str:
    """
    Extrae texto de la Responses API de forma robusta.
    """
    # SDKs recientes
    try:
        txt = (getattr(resp, "output_text", None) or "").strip()
        if txt:
            return txt
    except Exception:
        pass

    # Fallback por si cambia el SDK
    try:
        partes = []
        for item in getattr(resp, "output", []) or []:
            if getattr(item, "type", "") == "message":
                for c in getattr(item, "content", []) or []:
                    if getattr(c, "type", "") in ("output_text", "text"):
                        partes.append(getattr(c, "text", ""))
        return "".join(partes).strip()
    except Exception:
        return ""

def _chat_to_text(resp) -> str:
    """
    Extrae texto de Chat Completions (fallback).
    """
    try:
        choice = (resp.choices or [None])[0]
        msg = getattr(choice, "message", None)
        if msg and getattr(msg, "content", None):
            return (msg.content or "").strip()
    except Exception:
        pass
    return ""

def _chat_create_safe(**kw):
    """
    ENVÍA usando Responses si es posible; si no, cae a Chat Completions.
    Acepta: model, messages, max_completion_tokens / max_tokens / max_output_tokens,
            temperature, tools, tool_choice, metadata, input.
    Devuelve SIEMPRE el texto (str).
    """
    # Normalizar temperature
    if kw.get("temperature", None) is None:
        kw.pop("temperature", None)

    # Mapear topes de salida (no mutar `kw` original más de lo necesario)
    tok = kw.pop("max_output_tokens", None) or kw.pop("max_completion_tokens", None) or kw.pop("max_tokens", None)
    if tok is not None:
        try:
            tok = int(tok)
        except Exception:
            tok = None

    # Aceptamos `messages` o `input` (multimodal). Los pasamos intactos al router.
    messages = kw.get("messages", None)
    input_items = kw.get("input", None)
    model = kw.get("model", MODEL_ANALISIS)

    # Payload “neutro” para el router
    attempt = {"model": model}
    if messages is not None:
        attempt["messages"] = messages
    if input_items is not None:
        attempt["input"] = input_items
    if tok is not None:
        attempt["max_output_tokens"] = tok
    # Passthrough
    for k in ("temperature", "tools", "tool_choice", "metadata"):
        if k in kw:
            attempt[k] = kw[k]

    try:
        r = _responses_or_chat(client_timed, **attempt)
    except Exception as e:
        return f"[OPENAI-ERROR] {e}"

    # Extraer texto según tipo de respuesta
    # Heurística rápida: si tiene .output_text, es Responses
    txt = ""
    try:
        if hasattr(r, "output_text") or hasattr(r, "output"):
            txt = _resp_to_text(r)
        else:
            txt = _chat_to_text(r)
    except Exception:
        txt = ""

    return txt or ""

# ==================== Extracción por tipo de archivo ====================
def _leer_todo(file) -> bytes:
    try:
        file.file.seek(0)
        raw = file.file.read()
    except Exception:
        try:
            raw = file.read()
        except Exception:
            raw = b""
    return raw or b"""

def _ext_de_archivo(file) -> str:
    nombre = getattr(file, "filename", "") or ""
    _, ext = os.path.splitext(nombre)
    return (ext or "").lower().strip()

def _mime_guess(file) -> str:
    nombre = getattr(file, "filename", "") or ""
    m, _ = mimetypes.guess_type(nombre)
    return m or ""

def _texto_nativo_etiquetado(doc: fitz.Document) -> str:
    partes: List[str] = []
    for i, p in enumerate(doc, 1):
        t = (p.get_text() or "").strip()
        if t:
            partes.append(f"[PÁGINA {i}]\n{t}")
        else:
            partes.append(f"[PÁGINA {i}] (sin texto)")
    return "\n\n".join(partes).strip()

def extraer_texto_de_pdf(file) -> str:
    t0 = _t()
    raw = _leer_todo(file)
    if not raw:
        _log_tiempo("extraccion_pdf_sin_bytes", t0)
        return ""

    try:
        with fitz.open(stream=raw, filetype="pdf") as doc:
            suma = 0
            for p in doc:
                suma += len((p.get_text() or "").strip())

            if suma < 500:
                ocr_t0 = _t()
                # OCR selectivo (muestreado)
                n = len(doc)
                to_process = min(n, VISION_MAX_PAGES)
                if to_process >= n:
                    page_idxs = list(range(n))
                else:
                    page_idxs = sorted({
                        int(round(i * (n - 1) / max(1, to_process - 1)))
                        for i in range(to_process)
                    })

                def _proc_page(i: int) -> Tuple[int, str]:
                    p = doc.load_page(i)
                    txt_nat = (p.get_text() or "").strip()
                    if len(txt_nat) >= OCR_TEXT_MIN_CHARS:
                        return i, f"[PÁGINA {i+1}]\n{txt_nat}"
                    png_bytes = _rasterizar_pagina(p)
                    b64 = base64.b64encode(png_bytes).decode("utf-8")
                    # Vision con wrapper unificado
                    prompt = (
                        "Extraé el TEXTO literal de esta imagen escaneada de un pliego. "
                        "Conservá títulos, tablas como líneas con separadores, listas y números. No resumas ni interpretes."
                    )
                    txt = _chat_create_safe(
                        model=VISION_MODEL,
                        messages=[{
                            "role": "user",
                            "content": [
                                {"type": "input_text", "text": prompt},
                                {"type": "input_image", "image_url": {"url": f"data:image/png;base64,{b64}"}}
                            ]
                        }],
                        max_output_tokens=900,
                    )
                    return i, (f"[PÁGINA {i+1}]\n{txt}" if txt else f"[PÁGINA {i+1}] (sin texto OCR)")

                resultados_map: Dict[int, str] = {}
                from concurrent.futures import ThreadPoolExecutor, as_completed
                with ThreadPoolExecutor(max_workers=max(1, OCR_CONCURRENCY)) as ex:
                    futs = [ex.submit(_proc_page, i) for i in page_idxs]
                    for fut in as_completed(futs):
                        try:
                            i, s = fut.result()
                            resultados_map[i] = s
                        except Exception:
                            pass

                orden = sorted(resultados_map.keys())
                res = [resultados_map[i] for i in orden]

                if n > to_process:
                    res.append(f"\n[AVISO] OCR muestreó {to_process}/{n} páginas distribuidas.")

                out_ocr = "\n\n".join([r for r in res if r]).strip()
                _log_tiempo("ocr_selectivo", ocr_t0)
                _log_tiempo("extraccion_pdf_total", t0)
                return out_ocr

            out = (
                _texto_nativo_etiquetado(doc)
                if PAGINAR_TEXTO_NATIVO
                else "\n".join([(p.get_text() or "") for p in doc])
            )
            _log_tiempo("extraccion_pdf_total", t0)
            return (out or "").strip()
    except Exception:
        try:
            out = raw.decode("utf-8", errors="ignore")
            _log_tiempo("extraccion_pdf_decode", t0)
            return out
        except Exception:
            _log_tiempo("extraccion_pdf_error", t0)
            return ""

def extraer_texto_de_docx(file) -> str:
    t0 = _t()
    raw = _leer_todo(file)
    if not raw:
        _log_tiempo("extraccion_docx_sin_bytes", t0)
        return ""

    try:
        import docx  # python-docx (lazy)
    except Exception:
        try:
            out = raw.decode("utf-8", errors="ignore")
            _log_tiempo("extraccion_docx_decode", t0)
            return out
        except Exception:
            _log_tiempo("extraccion_docx_error", t0)
            return ""

    try:
        document = docx.Document(io.BytesIO(raw))
        partes: List[str] = []

        for p in document.paragraphs:
            txt = (p.text or "").strip()
            if txt:
                partes.append(txt)

        for tbl in document.tables:
            for row in tbl.rows:
                celdas = [(cell.text or "").strip() for cell in row.cells]
                fila = " | ".join([c for c in celdas if c is not None])
                if fila.strip():
                    partes.append(fila)

        out = "\n".join(partes).strip()
        _log_tiempo("extraccion_docx_total", t0)
        return out
    except Exception:
        try:
            out = raw.decode("utf-8", errors="ignore")
            _log_tiempo("extraccion_docx_decode_fallback", t0)
            return out
        except Exception:
            _log_tiempo("extraccion_docx_error", t0)
            return ""

def extraer_texto_de_imagen(file) -> str:
    t0 = _t()
    raw = _leer_todo(file)
    if not raw:
        _log_tiempo("extraccion_imagen_sin_bytes", t0)
        return ""

    ext = _ext_de_archivo(file)
    try:
        img_doc = fitz.open(stream=raw, filetype=ext.lstrip(".") or None)
        page = img_doc.load_page(0)
        png = page.get_pixmap(alpha=False).tobytes("png")
        b64 = base64.b64encode(png).decode("utf-8")
        mime = "image/png"
    except Exception:
        b64 = base64.b64encode(raw).decode("utf-8")
        mime = "image/png" if ext == ".png" else "image/jpeg"

    prompt = (
        "Extraé el TEXTO literal de esta imagen escaneada de un pliego. "
        "Conservá títulos, tablas como líneas con separadores, listas y números. No resumas ni interpretes."
    )
    out = _chat_create_safe(
        model=VISION_MODEL,
        messages=[{
            "role": "user",
            "content": [
                {"type": "input_text", "text": prompt},
                {"type": "input_image", "image_url": {"url": f"data:{mime};base64,{b64}"}}
            ]
        }],
        max_output_tokens=900,
    )
    _log_tiempo("extraccion_imagen_ocr", t0)
    return out or ""

def extraer_texto_universal(file) -> str:
    t0 = _t()
    ext = _ext_de_archivo(file)
    mime = _mime_guess(file)

    if ext == ".pdf" or (mime == "application/pdf"):
        out = extraer_texto_de_pdf(file)
        _log_tiempo("extraer_texto_universal_pdf", t0)
        return out

    if ext == ".docx" or (mime in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"]):
        out = extraer_texto_de_docx(file)
        _log_tiempo("extraer_texto_universal_docx", t0)
        return out

    if ext in [".png", ".jpg", ".jpeg", ".webp"] or (mime.startswith("image/") if mime else False):
        out = extraer_texto_de_imagen(file)
        _log_tiempo("extraer_texto_universal_imagen", t0)
        return out

    raw = _leer_todo(file)
    if not raw:
        _log_tiempo("extraer_texto_universal_sin_bytes", t0)
        return ""

    try:
        text = raw.decode("utf-8", errors="ignore")
    except Exception:
        text = ""

    if ext == ".rtf":
        text = re.sub(r"{\\rtf1.*?\\viewkind4\\uc1", "", text, flags=re.S)
        text = re.sub(r"\\[a-z]+-?\d* ?", "", text)
        text = text.replace("{", "").replace("}", "")

    out = (text or "").strip()
    _log_tiempo("extraer_texto_universal_texto_plano", t0)
    return out

# ==================== Pre-limpieza y helpers ====================
def _limpieza_basica_preanalisis(s: str) -> str:
    s = re.sub(r"\n?P[aá]gina\s+\d+\s+de\s+\d+\s*\n", "\n", s, flags=re.I)
    s = re.sub(r"\n[-_]{3,}\n", "\n", s)
    s = re.sub(r"[ \t]+\n", "\n", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return (s or "").strip()

_META_PATTERNS = [
    re.compile(r"(?i)\bparte\s+\d+\s+de\s+\d+"),
    re.compile(r"(?i)informe\s+basado\s+en\s+la\s+parte"),
    re.compile(r"(?i)revise\s+las\s+partes\s+restantes"),
    re.compile(r"(?i)informaci[oó]n\s+puede\s+estar\s+incompleta"),
    re.compile(r"(?i)^\s*informe\s+completo\s*$"),
    re.compile(r"(?i)^\s*informe\s+original\s*$"),
]

def _limpiar_meta(texto: str) -> str:
    lineas = []
    for ln in (texto or "").splitlines():
        if any(p.search(ln) for p in _META_PATTERNS):
            continue
        lineas.append(ln)
    return re.sub(r"\n{3,}", "\n\n", "\n".join(lineas)).strip()

def _particionar(texto: str, max_chars: int) -> List[str]:
    return [texto[i:i + max_chars] for i in range(0, len(texto or ""), max_chars)]
# -*- coding: utf-8 -*-
# utils.py — Parte 3/4 (Normalización + Hints + Secciones)

# --- Parche de alias por compatibilidad (si alguna parte usó __leer_todo) ---
try:
    __leer_todo  # type: ignore
except NameError:
    __leer_todo = _leer_todo  # alias seguro

# =============== Normalización de citas (multi vs único anexo) ===============
_CITA_ANEXO_RE = re.compile(r"\(Anexo\s+([IVXLCDM\d]+)(?:,\s*p\.\s*(\d+))?\)", re.I)

def _normalize_citas_salida(texto: str, varios_anexos: bool) -> str:
    if varios_anexos:
        return texto or ""
    # Si NO hay múltiples anexos, simplifica "(Anexo X, p. N)" => "(p. N)"
    def repl(m):
        pag = m.group(2)
        if pag:
            return f"(p. {pag})"
        return "(Fuente: documento provisto)"
    return _CITA_ANEXO_RE.sub(repl, texto or "")

# ==================== Normalización para PDF (sin markdown) ====================
_HDR_RE = re.compile(r"^\s{0,3}(#{1,6})\s*(.+)$")
_BULLET_RE = re.compile(r"^\s*[-*•]\s+")
_TABLE_SEP_RE = re.compile(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$")
_CODE_FENCE_RE = re.compile(r"^\s*```.*$")
_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
_BOLD_ITALIC_RE = re.compile(r"(\*\*|\*|__|_)(.*?)\1")

def _title_case(s: str) -> str:
    return " ".join(w.capitalize() if w else w for w in re.split(r"(\s+)", s or ""))

def preparar_texto_para_pdf(markdown_text: str) -> str:
    out_lines: List[str] = []
    for raw_ln in (markdown_text or "").splitlines():
        ln = raw_ln.rstrip()

        # Normaliza encabezado de Ficha si el modelo lo numeró como "0)"
        if re.match(r"^\s*0\)\s*Ficha\s+estandarizada\s+del\s+procedimiento\b", ln, flags=re.I):
            ln = "Ficha estandarizada del procedimiento (campos estandarizados)"

        if _CODE_FENCE_RE.match(ln):
            continue
        if re.match(r"(?i)^\s*informe\s+completo\s*$", ln):
            continue
        if re.match(r"(?i)^\s*informe\s+original\s*$", ln):
            continue

        m = _HDR_RE.match(ln)
        if m:
            titulo = _title_case(m.group(2).strip(": ").strip())
            out_lines.append(titulo)
            out_lines.append("")  # espacio tras título
            continue

        if _TABLE_SEP_RE.match(ln):
            continue

        if _BULLET_RE.match(ln):
            ln = _BULLET_RE.sub("• ", ln)

        ln = _LINK_RE.sub(lambda mm: f"{mm.group(1)} ({mm.group(2)})", ln)
        ln = _BOLD_ITALIC_RE.sub(lambda mm: mm.group(2), ln)
        out_lines.append(ln)

        if ln.strip().endswith(":"):
            out_lines.append("")  # espacio extra tras línea-título

    texto = "\n".join(out_lines)
    texto = re.sub(r"\n{3,}", "\n\n", texto).strip()
    return texto

# ==================== Hints regex (recall) ====================
def _buscar_candidatos(texto: str, pats: List[str], idx_pag: List[Tuple[int, int]], limit: int) -> List[str]:
    hits: List[str] = []
    for pat in pats:
        for m in re.finditer(pat, texto or "", flags=re.I):
            pos = m.start()
            p = _pagina_de_indice(idx_pag, pos)
            start = max(0, pos - 160)
            end = min(len(texto), pos + 240)
            snippet = (texto[start:end]).replace("\n", " ").strip()
            hits.append(f"- p. {p}: {snippet}")
            if len(hits) >= limit:
                return hits
    return hits[:limit]

def _build_regex_hints(texto: str, limit_per_field: Optional[int] = None, max_chars: Optional[int] = None) -> str:
    if not texto:
        return ""
    if limit_per_field is None:
        limit_per_field = HINTS_PER_FIELD
    if max_chars is None:
        max_chars = HINTS_MAX_CHARS

    idx_pag = _index_paginas(texto)
    secciones: List[str] = []
    for key, meta in DETECTABLE_FIELDS.items():
        hits = _buscar_candidatos(texto, meta["pats"], idx_pag, limit_per_field)
        if hits:
            secciones.append(f"[{meta['label']}]\n" + "\n".join(hits))
        if sum(len(s) for s in secciones) > max_chars:
            break
    return "\n\n".join(secciones[:])

DETECTABLE_FIELDS: Dict[str, Dict] = {
    "mant_oferta": {"label": "Mantenimiento de oferta", "pats": [r"mantenim[ií]ento de la oferta", r"validez de la oferta"]},
    "gar_mant":    {"label": "Garantía de mantenimiento", "pats": [r"garant[ií]a.*manten", r"\b5 ?%"]},
    "gar_cumpl":   {"label": "Garantía de cumplimiento", "pats": [r"garant[ií]a.*cumpl", r"\b10 ?%"]},
    "plazo_ent":   {"label": "Plazo de entrega", "pats": [r"plazo de entrega", r"\b\d{1,3}\s*d[ií]as"]},
    "tipo_cambio": {"label": "Tipo de cambio", "pats": [r"Banco\s+Naci[oó]n", r"tipo de cambio", r"\bBNA\b"]},
    "comision":    {"label": "Comisión de (Pre)?Adjudicación", "pats": [r"Comisi[oó]n.*(pre)?adjudicaci[oó]n"]},
    "muestras":    {"label": "Muestras", "pats": [r"\bmuestras?\b"]},
    "planilla":    {"label": "Planilla de cotización y renglones", "pats": [r"planilla.*cotizaci[oó]n", r"renglones?"]},
    "modalidad":   {"label": "Procedimiento/Modalidad", "pats": [r"licitaci[oó]n\s+(p[úu]blica|privada)", r"contrataci[oó]n\s+directa", r"compra\s+menor", r"subasta", r"modalidad"]},
    "plazo_contr": {"label": "Duración del contrato", "pats": [r"duraci[oó]n del contrato", r"plazo contractual", r"por el t[ée]rmino\s+de\s+\d+", r"\b\d{1,4}\s*d[ií]as"]},
    "prorroga":    {"label": "Prórroga/Ampliación", "pats": [r"pr[oó]rroga", r"ampliaci[oó]n", r"hasta\s+el\s+100%"]},
    "presupuesto": {"label": "Monto / Presupuesto", "pats": [r"presupuesto (estimado|oficial|referencial)", r"monto\s+estimado", r"cr[ée]dito\s+disponible", r"\$\s?\d{1,3}(?:\.\d{3})*(?:,\d{2})?"]},
    "expediente":  {"label": "Expediente / N° proceso", "pats": [r"\bEX-\d{4}-[A-Z0-9-]+", r"\bN[°º]\s*de\s*(proceso|procedimiento|expediente)"]},
    "fechas":      {"label": "Fechas y horas", "pats": [r"\b\d{2}/\d{2}/\d{4}\b", r"\b\d{1,2}:\d{2}\s*(?:hs|h)\b"]},
    "contacto":    {"label": "Contactos y portales", "pats": [r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", r"https?://[^\s)]+|www\.[^\s)]+"]},
    "costo_pliego":{"label": "Costo/valor del pliego", "pats": [r"(?:costo|valor)\s+del\s+pliego", r"adquisici[oó]n\s+del\s+pliego", r"\$\s?\d{1,3}(?:\.\d{3})*(?:,\d{2})?"]},
    "subsanacion": {"label": "Subsanación", "pats": [r"subsanaci[oó]n"]},
    "perf_modif":  {"label": "Perfeccionamiento/Modificaciones", "pats": [r"perfeccionamiento", r"modificaci[oó]n"]},
    "preferencias":{"label": "Preferencias", "pats": [r"preferencias"]},
    "criterios":   {"label": "Criterios de evaluación", "pats": [r"criterios?\s+de\s+evaluaci[oó]n"]},
    "renglones":   {"label": "Renglones y especificaciones", "pats": [r"Rengl[oó]n\s*\d+", r"Especificaciones?\s+t[ée]cnicas?"]},
    "articulos":   {"label": "Artículos citados", "pats": [r"\bArt(?:[íi]culo|\.)\s*\d+[A-Za-z]?\b"]},
    "estado":      {"label": "Estado del trámite", "pats": [r"\bestado\b", r"\bvigente\b", r"\b(adjudicado|desierto|fracasado|cerrado)\b"]},
    "consultas":   {"label": "Inicio y final de consultas", "pats": [r"\bconsultas\b", r"aclaraciones", r"preguntas"]},
    "apertura":    {"label": "Acto de apertura", "pats": [r"acto\s+de\s+apertura", r"\bapertura\b"]},
    "tipo_cotiz":  {"label": "Tipo de cotización", "pats": [r"forma\s+de\s+cotizaci[oó]n", r"tipo\s+de\s+cotizaci[oó]n", r"cotizaci[oó]n\s+por"]},
    "tipo_adj":    {"label": "Tipo de adjudicación", "pats": [r"adjudicaci[oó]n\s+por\s+(rengl[oó]n|lote|total)"]},
    "moneda":      {"label": "Moneda", "pats": [r"\bmoneda\b", r"\bARS\b", r"\bUSD\b"]},
    "obj_gasto":   {"label": "Objeto del gasto", "pats": [r"objeto\s+del\s+gasto", r"partida\s+presupuestaria", r"clasificador"]},
    "ofertas_perm":{"label": "Ofertas permitidas", "pats": [r"m[aá]s\s+de\s+una\s+oferta", r"ofertas?\s+alternativas", r"una\s+sola\s+oferta"]},
}

# ==================== Utilidades de conteo y evidencia ====================
def _count(pattern: str, text: str) -> int:
    return len(re.findall(pattern, text or "", flags=re.I))

_ART_HEAD_RE = re.compile(r"(?im)^\s*(art(?:[íi]culo|\.?)\s*\d+[a-zº°]?)\s*[-–—:]?\s*(.*)$")
_ART_BLOCK_RE = re.compile(
    r"(?ims)^\s*(art(?:[íi]culo|\.?)\s*\d+[a-zº°]?)\s*[-–—:]?\s*(.+?)(?=^\s*art(?:[íi]culo|\.?)\s*\d+[a-zº°]?|\Z)"
)

def _extraer_articulos_con_snippets(texto: str) -> List[Tuple[str, str, int, Optional[int]]]:
    texto = texto or ""
    idx = _index_paginas(texto)
    idx_ax = _index_anexos(texto)
    res: List[Tuple[str, str, int, Optional[int]]] = []

    for m in _ART_BLOCK_RE.finditer(texto):
        start = m.start()
        p = _pagina_de_indice(idx, start)
        ax = _anexo_en_pos(idx_ax, start)
        rotulo = (m.group(1) or "").strip()
        contenido = (m.group(2) or "").strip()
        snippet = contenido[:200].replace("\n", " ").strip()
        res.append((rotulo, snippet, p, ax))

    if not res:
        for m in _ART_HEAD_RE.finditer(texto):
            start = m.start()
            p = _pagina_de_indice(idx, start)
            ax = _anexo_en_pos(idx_ax, start)
            rotulo = (m.group(1) or "").strip()
            contenido = ((m.group(2) or "").strip())
            snippet = (contenido[:200]).replace("\n", " ").strip()
            res.append((rotulo, snippet, p, ax))

    return res

_ROW_START_RE = re.compile(r"(?im)^(?:reng(?:l[oó]n)?\.?\s*)(\d{1,4})\b")
_CODE_RE = re.compile(r"\b[A-Z]{1,3}\d{5,8}\b")
_QTY_RE  = re.compile(r"\b\d{1,6}\b")

def _extraer_renglones_y_especificaciones(texto: str) -> List[Tuple[int, Optional[int], Optional[str], str, int, Optional[int]]]:
    texto = texto or ""
    idx = _index_paginas(texto)
    idx_ax = _index_anexos(texto)
    res: List[Tuple[int, Optional[int], Optional[str], str, int, Optional[int]]] = []

    lines = texto.splitlines()
    pos = 0
    starts: List[Tuple[int, int]] = []
    for i, ln in enumerate(lines):
        m = _ROW_START_RE.match(ln)
        if m:
            starts.append((i, pos))
        pos += len(ln) + 1

    if not starts:
        return res

    starts.append((len(lines), len(texto)))

    for k in range(len(starts) - 1):
        i_line, abs_pos = starts[k]
        j_line, _abs_pos_next = starts[k + 1]
        block_lines = lines[i_line:j_line]
        block_text = " ".join([re.sub(r"\s+", " ", x).strip() for x in block_lines if x.strip()])

        mnum = _ROW_START_RE.match(lines[i_line])
        try:
            num_r = int(mnum.group(1)) if mnum else None
        except Exception:
            num_r = None

        qty = None
        if mnum:
            tail = lines[i_line][mnum.end():]
            mqty = _QTY_RE.search(tail)
            if mqty:
                try:
                    qty = int(mqty.group(0))
                except Exception:
                    qty = None

        mcode = _CODE_RE.search(block_text)
        code = mcode.group(0) if mcode else None

        desc = block_text
        if code:
            desc = re.sub(re.escape(code), "", desc)
        if qty is not None:
            desc = re.sub(rf"\b{qty}\b", "", desc)
        if num_r is not None:
            desc = re.sub(rf"^\s*{num_r}\b", "", desc)
        desc = re.sub(r"\s+", " ", desc).strip()

        p = _pagina_de_indice(idx, abs_pos)
        ax = _anexo_en_pos(idx_ax, abs_pos)

        if num_r is not None:
            res.append((num_r, qty, code, desc, p, ax))

    res.sort(key=lambda t: t[0])
    return res

# ==================== Construcción de secciones (2.13, 2.16, 2.3, 2.15) ====================
_ART_KEYS = re.compile(
    r"(objeto|tipolog|modalidad|mantenim|pr[oó]rroga|oferta|apertura|evaluaci[oó]n|empate|mejora|adjudicaci[oó]n|"
    r"garant[ií]a|entrega|plazo|pago|factura|sanci[oó]n|penalidad|rescis[ií]n|perfeccionamiento|subsanaci[oó]n)",
    re.I
)

def _truncate_words(s: str, max_words: int) -> str:
    try:
        words = re.findall(r"\S+", s or "")
        if len(words) <= max_words:
            return (s or "").strip()
        return " ".join(words[:max_words]).rstrip(",.;:") + "..."
    except Exception:
        return (s or "").strip()

def _build_section_213(texto: str, varios_anexos: bool) -> str:
    rows = _extraer_renglones_y_especificaciones(texto or "")
    if not rows:
        return ""
    rows = rows[:max(1, MAX_RENGLONES_OUT)]
    lines = ["2.13 Planilla de cotización y renglones:"]
    for (num, qty, code, desc, p, ax) in rows:
        desc_corta = _truncate_words(desc, RENGLON_DESC_MAX_WORDS)
        partes = [f"Renglón {num}"]
        if qty is not None:
            partes.append(f"Cantidad: {qty}")
        if code:
            partes.append(f"Código: {code}")
        partes.append(f"Descripción/Especificaciones: {desc_corta}")
        cita = f"(Anexo {ax}, p. {p})" if varios_anexos and ax else (f"(p. {p})" if p else "(Fuente: documento provisto)")
        lines.append(" - " + " — ".join(partes) + f" {cita}")
    return "\n".join(lines)

def _build_section_216(texto: str, varios_anexos: bool) -> str:
    arts = _extraer_articulos_con_snippets(texto or "")
    if not arts:
        return ""
    arts = [(rot, sn, p, ax) for (rot, sn, p, ax) in arts if _ART_KEYS.search(sn or "") or _ART_KEYS.search(rot or "")]
    if not arts:
        return ""
    arts = arts[:max(1, MAX_ARTICULOS_OUT)]
    lines = ["2.16 Catálogo de artículos citados:"]
    for (rot, sn, p, ax) in arts:
        rot_norm = re.sub(r"(?i)art(?:[íi]culo|\.)\s*", "Art. ", rot or "").strip()
        sn = _truncate_words(sn or "", ART_SNIPPET_MAX_WORDS)
        cita = f"(Anexo {ax}, p. {p})" if varios_anexos and ax else (f"(p. {p})" if p else "(Fuente: documento provisto)")
        lines.append(f" - {rot_norm} — {sn} {cita}")
    return "\n".join(lines)

CONTACT_EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
CONTACT_URL_RE   = re.compile(r"(https?://[^\s)]+|www\.[^\s)]+)")

def _extraer_contactos_con_paginas(texto: str) -> List[Tuple[str, str, int, Optional[int]]]:
    texto = texto or ""
    idx_pag = _index_paginas(texto)
    idx_ax  = _index_anexos(texto)
    res: List[Tuple[str, str, int, Optional[int]]] = []

    for m in CONTACT_EMAIL_RE.finditer(texto):
        pos = m.start()
        p = _pagina_de_indice(idx_pag, pos)
        ax = _anexo_en_pos(idx_ax, pos)
        res.append(("email", m.group(0), p, ax))

    for m in CONTACT_URL_RE.finditer(texto):
        pos = m.start()
        p = _pagina_de_indice(idx_pag, pos)
        ax = _anexo_en_pos(idx_ax, pos)
        v = (m.group(0) or "").rstrip(").,;")
        res.append(("url", v, p, ax))

    seen = set()
    dedup: List[Tuple[str, str, int, Optional[int]]] = []
    for t, v, p, ax in res:
        key = (t, v.lower())
        if key in seen:
            continue
        seen.add(key)
        dedup.append((t, v, p, ax))
    return dedup

def _build_section_23(texto: str, varios_anexos: bool) -> str:
    items = _extraer_contactos_con_paginas(texto or "")
    if not items:
        return ""
    out = ["2.3 Contactos y portales:"]
    for (t, v, p, ax) in items:
        etiqueta = "Email" if t == "email" else "URL"
        cita = f"(Anexo {ax}, p. {p})" if varios_anexos and ax else (f"(p. {p})" if p else "(Fuente: documento provisto)")
        out.append(f" - {etiqueta}: {v} {cita}")
    return "\n".join(out)

NORM_TIPOS = [
    ("Ley",        r"\bLey(?:\s*N[°º])?\s*([\d\.]{1,7}(?:/\d{2,4})?)\b"),
    ("Decreto",    r"\bDecreto(?:\s*N[°º])?\s*([\d\.]{1,7}(?:/\d{2,4})?)\b"),
    ("Resolución", r"\bResoluci[oó]n(?:\s*(?:Ministerial|Conjunta))?\s*(?:N[°º]\s*)?(\d{1,7}(?:/\d{2,4})?)\b"),
    ("Disposición",r"\bDisposici[oó]n\s*(?:N[°º]\s*)?(\d{1,7}(?:/\d{2,4})?)\b"),
]
NORM_PATTS = [(tipo, re.compile(patt, re.I)) for (tipo, patt) in NORM_TIPOS]

def _extraer_normativa(texto: str) -> List[Tuple[str, str, int, Optional[int]]]:
    texto = texto or ""
    idx_pag = _index_paginas(texto)
    idx_ax  = _index_anexos(texto)
    res: List[Tuple[str, str, int, Optional[int]]] = []

    for (tipo, cre) in NORM_PATTS:
        for m in cre.finditer(texto):
            pos = m.start()
            p = _pagina_de_indice(idx_pag, pos)
            ax = _anexo_en_pos(idx_ax, pos)
            numero = (m.group(1) or "").strip()
            res.append((tipo, numero, p, ax))

    seen = set()
    dedup: List[Tuple[str, str, int, Optional[int]]] = []
    for t, n, p, ax in res:
        key = (t.lower(), n)
        if key in seen:
            continue
        seen.add(key)
        dedup.append((t, n, p, ax))
    return dedup

def _build_section_215(texto: str, varios_anexos: bool) -> str:
    normas = _extraer_normativa(texto or "")
    if not normas:
        return ""
    out = ["2.15 Normativa aplicable:"]
    for (t, n, p, ax) in normas:
        cita = f"(Anexo {ax}, p. {p})" if varios_anexos and ax else (f"(p. {p})" if p else "(Fuente: documento provisto)")
        out.append(f" - {t} {n} {cita}")
    return "\n".join(out)

# ====== Reemplazo/inyección de secciones en el informe ======
def _find_section_bounds(text: str, header_regex: str) -> Tuple[int, int]:
    text = text or ""
    m = re.search(header_regex, text, flags=re.I)
    if not m:
        return (-1, -1)
    start = m.start()
    nxt = re.search(r"(?im)^\s*2\.(1[0-9]|[1-9])\s", text[m.end():])
    if not nxt:
        return (start, len(text))
    return (start, m.end() + nxt.start())

def _replace_section(text: str, header_regex: str, replacement: str) -> str:
    s, e = _find_section_bounds(text or "", header_regex)
    if s == -1:
        return (text or "").rstrip() + "\n\n" + (replacement or "").strip() + "\n"
    return (text or "")[:s] + (replacement or "").strip() + "\n" + (text or "")[e:]

def _ampliar_secciones_especificas(informe: str, texto_fuente: str, varios_anexos: bool) -> str:
    out = informe or ""

    # Siempre normalizar 2.3 y 2.15 desde extracción determinística (evita ruido/omisiones)
    sec23 = _build_section_23(texto_fuente or "", varios_anexos)
    if sec23:
        out = _replace_section(out, r"(?im)^\s*2\.3\s+Contactos", sec23)

    sec215 = _build_section_215(texto_fuente or "", varios_anexos)
    if sec215:
        out = _replace_section(out, r"(?im)^\s*2\.15\s+Normativa", sec215)

    if not EXPAND_SECTIONS_213_216:
        return out

    sec213 = _build_section_213(texto_fuente or "", varios_anexos)
    if sec213:
        alt213 = sec213.replace("2.13 Planilla de cotización y renglones:", "9) Renglones y planilla de cotización:")
        out = _replace_section(out, r"(?im)^\s*9\)\s*Renglones\s+y\s+planilla", alt213)
        out = _replace_section(out, r"(?im)^\s*2\.13\s+Planilla", sec213)

    sec216 = _build_section_216(texto_fuente or "", varios_anexos)
    if sec216:
        out = _replace_section(out, r"(?im)^\s*2\.16\s+Cat[áa]logo\s+de\s+art", sec216)
        # remueve posibles encabezados redundantes generados por el modelo
        out = re.sub(r"(?im)^\s*(ANEXO|Anexo)\s*[-–—]?\s*Cat[áa]logo\s+de\s+art[^\n]*\n?", "", out)

    out = re.sub(r"(?im)^\s*informe\s+original\s*$", "", out)
    return out

# === Post-procesos de Ficha y secciones ===
def _reparar_ficha(informe: str, texto_fuente: str) -> str:
    try:
        total_renglones = len(_extraer_renglones_y_especificaciones(texto_fuente or ""))
    except Exception:
        total_renglones = 0

    if total_renglones:
        informe = re.sub(
            r"(?im)^(\s*•\s*(?:N[uú]mero\s+de\s+rengl[oó]n|Numero\s+de\s+renglon)\s*:\s*)[^\n]*$",
            lambda m: f"{m.group(1)}Total de renglones: {total_renglones}; ver Sección 9 para el detalle completo",
            informe or ""
        )
        informe = re.sub(
            r"(?im)\bTotal de renglones:\s*N\b",
            f"Total de renglones: {total_renglones}",
            informe or ""
        )

    informe = re.sub(
        r"(?im)^(\s*•\s*Monto:\s*)(?:\$+\s*\.{0,3}|[$…]+)\s*(\(.*?\))?\s*$",
        lambda m: f"{m.group(1)}NO ESPECIFICADO{(' ' + m.group(2) if m.group(2) else '')}",
        informe or ""
    )
    return (informe or "")

def _normalizar_encabezados_salida(informe: str) -> str:
    s = informe or ""
    s = re.sub(
        r"(?im)^\s*0\)\s*Ficha\s+estandarizada\s+del\s+procedimiento\s*\(campos\s+estandarizados\)\s*$",
        "Ficha estandarizada del procedimiento (campos estandarizados)", s
    )
    s = re.sub(
        r"(?im)^\s*0\)\s*Ficha\s+estandarizada\s+del\s+procedimiento\s*$",
        "Ficha estandarizada del procedimiento", s
    )
    return s.strip()

def _corregir_seccion_9_si_vacia(informe: str, texto_fuente: str, varios_anexos: bool) -> str:
    s = informe or ""
    start, end = _find_section_bounds(s, r"(?im)^\s*9\)\s*Renglones\s+y\s+planilla")
    needs_fix = False

    if start == -1:
        needs_fix = True
    else:
        bloque = s[start:end]
        if (_count(r"(?im)\bRengl[oó]n\s+\d+", bloque) == 0) or \
           (re.search(r"(?i)\bNO ESPECIFICADO\b", bloque) is not None) or \
           (len(bloque.strip()) < 80):
            needs_fix = True

    if not needs_fix:
        return s

    sec213 = _build_section_213(texto_fuente or "", varios_anexos)
    if not sec213:
        return s

    sec9 = sec213.replace("2.13 Planilla de cotización y renglones:", "9) Renglones y planilla de cotización:")
    if start == -1:
        return (s.rstrip() + "\n\n" + sec9.strip() + "\n")
    else:
        return _replace_section(s, r"(?im)^\s*9\)\s*Renglones\s+y\s+planilla", sec9)

# ====== Política de salida (guía para el modelo + recorte determinístico) ======
def _output_policy_block() -> str:
    return (
        "\n\n=== POLÍTICA DE LONGITUD Y ORDEN ===\n"
        f"- El informe completo NO debe exceder ~{MAX_WORDS_TOTAL_GUIDE} palabras.\n"
        f"- Cada sección 1–12 máximo {MAX_LINES_PER_SECTION} líneas; cada bullet/ítem máximo {MAX_WORDS_PER_BULLET} palabras.\n"
        f"- Evitar repeticiones; un solo dato por línea con su cita.\n"
        "- Prohibido agregar anexos, apéndices, “hallazgos” o texto fuera de la estructura pedida.\n"
        "- Mantener exactamente el orden: Ficha, 1) … 12)."
    )

def _split_informe_por_secciones(s: str) -> List[Tuple[str, str]]:
    """
    Devuelve lista [(header, body)] incluyendo "Ficha..." como primera si existe.
    """
    s = s or ""
    s = re.sub(r"\n{3,}", "\n\n", s)

    blocks: List[Tuple[str, str]] = []
    ficha_re = re.compile(r"(?im)^\s*Ficha\s+estandarizada\s+del\s+procedimiento.*$")
    sec_re = re.compile(r"(?im)^\s*(\d{1,2})\)\s")

    m_ficha = ficha_re.search(s)
    starts: List[Tuple[int, str]] = []

    if m_ficha:
        starts.append((m_ficha.start(), m_ficha.group(0).strip()))

    for m in sec_re.finditer(s):
        line_start = s.rfind("\n", 0, m.start()) + 1
        line_end = s.find("\n", m.start())
        if line_end == -1:
            line_end = len(s)
        header_line = s[line_start:line_end].strip()
        starts.append((line_start, header_line))

    if not starts:
        return [("INFORME", s.strip())]

    starts.sort(key=lambda t: t[0])
    starts.append((len(s), ""))  # centinela

    for i in range(len(starts) - 1):
        hpos, header = starts[i]
        npos, _ = starts[i+1]
        body = s[hpos: npos].split("\n", 1)
        if len(body) == 1:
            header_line, body_content = header, ""
        else:
            header_line, body_content = body[0].strip(), body[1].strip()
        blocks.append((header_line, body_content))

    return blocks

def _recortar_por_politica(header: str, body: str) -> str:
    """
    Aplica recortes determinísticos a un bloque de sección (body), manteniendo el header.
    """
    if not STRICT_OUT:
        return f"{header}\n{body}".strip()

    body = re.sub(r"(?im)^===.*$", "", body)

    lines = [ln for ln in (body or "").splitlines()]
    recortadas: List[str] = []
    max_lines = max(1, MAX_LINES_PER_SECTION)

    for ln in lines:
        ln_clean = ln.strip()
        if not ln_clean:
            recortadas.append("")
            continue
        ln_clean = _truncate_words(ln_clean, MAX_WORDS_PER_BULLET)
        recortadas.append(ln_clean)
        if len([x for x in recortadas if x.strip()]) >= max_lines:
            break

    body_recortado = "\n".join(recortadas).strip()

    if len(body_recortado) > SECTION_CHAR_LIMIT:
        body_recortado = body_recortado[:SECTION_CHAR_LIMIT].rsplit("\n", 1)[0].rstrip() + "\n..."

    return f"{header}\n{body_recortado}".strip()

def _enforce_output_policy(informe: str) -> str:
    """
    Aplica orden y recortes por sección; limita tamaño total.
    """
    s = informe or ""
    blocks = _split_informe_por_secciones(s)

    ficha = [b for b in blocks if re.match(r"(?im)^Ficha\s+estandarizada\s+del\s+procedimiento", b[0])]
    otros = [b for b in blocks if not re.match(r"(?im)^Ficha\s+estandarizada\s+del\s+procedimiento", b[0])]

    def _sec_num(h: str) -> int:
        m = re.match(r"^\s*(\d{1,2})\)\s", h)
        return int(m.group(1)) if m else 99

    otros.sort(key=lambda b: _sec_num(b[0]))
    ordered = ficha + otros

    partes: List[str] = []
    for (h, body) in ordered:
        partes.append(_recortar_por_politica(h, body))

    out = "\n\n".join([p for p in partes if p.strip()])

    if len(out) > MAX_TOTAL_CHARS_OUT:
        out = out[:MAX_TOTAL_CHARS_OUT].rsplit("\n", 1)[0].rstrip() + "\n..."

    out = re.sub(r"\n{3,}", "\n\n", out).strip()
    return out
# -*- coding: utf-8 -*-
# utils.py — Parte 4/4 (Pipeline de análisis + PDF + helpers finales + RAG opcional)

# ==================== Mensajería al modelo ====================
def _mk_temperature() -> Optional[float]:
    try:
        return float(TEMPERATURE_ANALISIS) if TEMPERATURE_ANALISIS else None
    except Exception:
        return None

def _pick_model(final_pass: bool = False) -> str:
    """
    Selecciona el modelo según modo/flags. Para pasadas parciales puede forzar uno 'fast'.
    """
    if not final_pass and (ANALISIS_MODO == "fast") and FAST_FORCE_MODEL:
        return FAST_FORCE_MODEL
    return MODEL_ANALISIS

def _craft_system_prompt(varios_anexos: bool, texto_hints: str = "", kb_context: str = "") -> str:
    """
    Arma el prompt de sistema con reglas, sinónimos, política de salida y opcionalmente
    un bloque de contexto KB (si viene no-vacío).
    """
    base = prompt_andres(varios_anexos)
    bloques = [base]

    sinos = _sinonimos_text()
    if sinos:
        bloques.append(sinos)

    # Nota sobre KB: solo guía terminológica. No se deben introducir datos inexistentes.
    if kb_context.strip():
        kb_safe = (
            "Usá el siguiente contexto SOLO para orientar terminología/búsqueda interna; "
            "NO cites ni agregues datos que NO estén en los archivos analizados.\n"
            + kb_context.strip()
        )
        bloques.append("\n=== CONTEXTO KB (referencia no-citable) ===\n" + kb_safe)

    bloques.append(_output_policy_block())

    if (texto_hints or "").strip():
        bloques.append("\n=== HINTS DETECTADOS (útiles para recall) ===\n" + texto_hints.strip())

    return "\n\n".join(b for b in bloques if b).strip()

def _msg_single_block(varios_anexos: bool, texto_fuente: str, texto_hints: str = "", titulo: str = "", kb_context: str = "") -> List[Dict[str, Any]]:
    sys = _craft_system_prompt(varios_anexos, texto_hints=texto_hints, kb_context=kb_context)
    user = []
    if titulo:
        user.append(f"TÍTULO/BLOQUE: {titulo}")
    user.append("CONTENIDO A ANALIZAR (texto literal paginado):")
    user.append((texto_fuente or "").strip())
    content = "\n\n".join(user)
    return [{"role": "system", "content": sys}, {"role": "user", "content": content}]

def _call_chat(messages: List[Dict[str, Any]], model: Optional[str] = None, max_tokens: Optional[int] = None) -> str:
    payload: Dict[str, Any] = {
        "model": model or _pick_model(final_pass=False),
        "messages": messages,
    }
    temp = _mk_temperature()
    if temp is not None:
        payload["temperature"] = temp
    if max_tokens is not None:
        # Responses API usa max_output_tokens
        payload["max_output_tokens"] = int(max_tokens)

    # _chat_create_safe (Parte 2) devuelve texto directo (string)
    return _chat_create_safe(**payload)

# ==================== RAG liviano sobre KB (sin NumPy) ====================
def _cosine_py(a: List[float], b: List[float]) -> float:
    # Cosine similarity sin dependencias externas
    if not a or not b:
        return 0.0
    n = min(len(a), len(b))
    dot = 0.0
    na = 0.0
    nb = 0.0
    for i in range(n):
        x = float(a[i])
        y = float(b[i])
        dot += x * y
        na += x * x
        nb += y * y
    na = (na ** 0.5) or 1e-8
    nb = (nb ** 0.5) or 1e-8
    return float(dot / (na * nb))

def _kb_try_open_session():
    """
    Intenta abrir una sesión SQLAlchemy del proyecto (database.SessionLocal).
    Si no existe, devuelve None y el RAG se desactiva silenciosamente.
    """
    try:
        from database import SessionLocal as _SessionLocal
        return _SessionLocal()
    except Exception:
        return None

def _kb_build_context_from_db(query: str, source_slug: Optional[str] = None, top_k: int = 8) -> str:
    """
    Recupera chunks de KB similares a la consulta y arma un bloque de contexto legible.
    Usa las tablas del proyecto (models.KBChunk/KBFile/KBSource) que ya se usan en la ingesta.
    """
    if not (query or "").strip():
        return ""

    # Embedding de la consulta
    try:
        q_emb = _kb_embed(query)
    except Exception:
        return ""

    db = _kb_try_open_session()
    if not db:
        return ""  # sin DB, sin contexto

    KBSource, KBFile, KBChunk, KBPriority = _kb_models()
    try:
        # Limitar por performance
        N = int(os.getenv("KB_MAX_CHUNKS_SEARCH", "4000"))

        q = db.query(KBChunk).order_by(KBChunk.id.desc()).limit(N)
        if source_slug:
            # filtrar por fuente si se pide
            src = db.query(KBSource).filter(KBSource.name == source_slug).first()
            if src:
                files_ids = [f.id for f in db.query(KBFile.id).filter(KBFile.source_id == src.id).all()]
                if files_ids:
                    q = q.filter(KBChunk.file_id.in_(files_ids))
        rows = list(q.all())[::-1]  # ascendente

        # Cargar prioridades (si existen)
        try:
            gprio = {r.label.lower(): float(r.weight) for r in db.query(KBPriority).all()}  # type: ignore[attr-defined]
        except Exception:
            gprio = {}

        scored: List[Tuple[float, Any]] = []
        for r in rows:
            try:
                emb = json.loads(r.embedding)  # en la ingesta guardamos json.dumps(vec)
            except Exception:
                continue
            score = _cosine_py(q_emb, emb)

            # BONUS por prioridades si algún término aparece
            txt_low = (r.text or "").lower()
            bonus = 0.0
            for t, w in gprio.items():
                if t in txt_low:
                    bonus += 0.03 * max(1.0, float(w))
            score += bonus

            scored.append((score, r))

        scored.sort(key=lambda t: t[0], reverse=True)
        if not scored:
            return ""

        lines = ["[KB] Extractos relevantes (NO citables — solo orientación):"]
        take = scored[:max(1, int(top_k))]
        for s, r in take:
            try:
                f = db.query(KBFile).filter(KBFile.id == r.file_id).first()
                src = db.query(KBSource).filter(KBSource.id == f.source_id).first() if f else None
                head = f"Fuente: {src.name if src else '-'} | Archivo: {getattr(f, 'filename', getattr(f, 'path', ''))} | Chunk #{getattr(r, 'ord', getattr(r, 'ordinal', 0))} | score={round(float(s), 4)}"
            except Exception:
                head = f"Chunk | score={round(float(s), 4)}"
            body = (r.text or "").strip()
            lines.append(f"--- {head}\n{body}\n")
        return "\n".join(lines).strip()
    except Exception:
        return ""
    finally:
        try:
            db.close()
        except Exception:
            pass

def build_kb_context(query: str, source: Optional[str] = None, top_k: int = 8) -> str:
    """
    Wrapper público: si hay DB y KB cargada, devuelve un bloque de contexto;
    si no, devuelve cadena vacía (el pipeline lo manejará).
    """
    try:
        return _kb_build_context_from_db(query=query, source_slug=source, top_k=top_k) or ""
    except Exception:
        return ""

# ==================== Análisis: single/multi-pass (con KB opcional) ====================
def _resumen_parcial(chunk_text: str, varios_anexos: bool, idx: int, total: int, texto_hints: str = "", kb_context: str = "") -> str:
    """
    Produce un resumen estructurado (mini-informe) del bloque.
    """
    titulo = f"Bloque {idx}/{total}"
    msgs = _msg_single_block(varios_anexos, chunk_text, texto_hints=texto_hints, titulo=titulo, kb_context=kb_context)
    out = _call_chat(msgs, model=_pick_model(final_pass=False), max_tokens=NOTAS_MAX_TOKENS)
    return out

def _agregar_y_consolidar(parciales: List[str], varios_anexos: bool, texto_hints: str = "", kb_context: str = "") -> str:
    """
    Funde los parciales en un informe único, aplicando la guía de salida.
    """
    sys = _craft_system_prompt(varios_anexos, texto_hints=texto_hints, kb_context=kb_context)
    corpus = "\n\n".join([f"=== PARCIAL {i+1} ===\n{p}" for i, p in enumerate(parciales) if (p or "").strip()])
    user = (
        "Integrá TODOS los parciales anteriores en un ÚNICO informe final, sin repetir texto, "
        "llenando las secciones que falten y citando correctamente (ver reglas). "
        "No agregues anexos ni texto fuera de la estructura pedida."
        "\n\n"
        + corpus
    )
    msgs = [{"role": "system", "content": sys}, {"role": "user", "content": user}]
    out = _call_chat(msgs, model=_pick_model(final_pass=True), max_tokens=MAX_COMPLETION_TOKENS_SALIDA)
    return out

def _postproceso_final(informe: str, texto_fuente: str, varios_anexos: bool) -> str:
    """
    Post-procesos determinísticos/seguros, manteniendo citas y estructura.
    """
    s = informe or ""
    s = _normalizar_encabezados_salida(s)
    s = _reparar_ficha(s, texto_fuente or "")
    s = _ampliar_secciones_especificas(s, texto_fuente or "", varios_anexos)
    s = _corregir_seccion_9_si_vacia(s, texto_fuente or "", varios_anexos)
    s = _normalize_citas_salida(s, varios_anexos)
    s = _enforce_output_policy(s)
    s = _limpiar_meta(s)
    s = re.sub(r"\n{3,}", "\n\n", s).strip()
    return s

def _input_incompleto(texto_fuente: str) -> bool:
    """
    Heurística para detectar texto incompleto/truncado:
    - OCR muestreado (agregado por extraer_texto_de_pdf)
    - Páginas sin texto OCR
    - Vacío o extremadamente corto
    """
    s = (texto_fuente or "").strip()
    if not s or len(s) < 40:
        return True
    if "[AVISO] OCR muestreó" in s:
        return True
    if "(sin texto OCR)" in s:
        return True
    return False

def analizar_y_generar_informe(
    texto_fuente: str,
    *,
    varios_anexos: Optional[bool] = None,
    force_multi: Optional[bool] = None,
    prefer_source: Optional[str] = None,
) -> str:
    """
    Pipeline principal para obtener el informe desde el texto crudo (paginado).
    - Limpia/normaliza.
    - Arma un bloque de CONTEXTO KB (si hay KB).
    - Opcionalmente genera HINTS regex para recall.
    - Decide single vs multi-pass y consolida.
    - Aplica post-procesos determinísticos.
    """
    t0 = _t()

    # Limpieza previa
    raw = (texto_fuente or "").strip()
    raw = _limpieza_basica_preanalisis(raw)
    raw = _limpiar_meta(raw)

    # Regla de “ERROR: texto de entrada incompleto”
    if _input_incompleto(raw):
        return "ERROR: texto de entrada incompleto"

    # Heurística de anexos si no se especifica
    if varios_anexos is None:
        varios_anexos = (_contar_anexos(raw) > 1)

    # Hints regex (opcionales)
    hints = _build_regex_hints(raw) if ENABLE_REGEX_HINTS else ""

    # Contexto KB (opcional y silencioso si no hay DB/KB)
    kb_query = raw[:1500]
    kb_ctx = build_kb_context(query=kb_query, source=prefer_source, top_k=10)

    # Elección single vs multi
    multi = bool(force_multi) or (len(raw) > MAX_SINGLE_PASS_CHARS)

    # Single-pass
    if not multi:
        msgs = _msg_single_block(
            varios_anexos,
            raw,
            texto_hints=hints,
            kb_context=kb_ctx
        )
        borrador = _call_chat(
            msgs,
            model=_pick_model(final_pass=True),
            max_tokens=MAX_COMPLETION_TOKENS_SALIDA
        )
        final = _postproceso_final(borrador, raw, varios_anexos)
        _log_tiempo("pipeline_single_pass", t0)
        return final

    # Multi-pass (parciales en paralelo)
    partes = _particionar(raw, max_chars=CHUNK_SIZE_BASE)
    parciales: List[str] = []

    def _work(i_chunk: int, total: int, texto: str) -> str:
        return _resumen_parcial(
            texto,
            varios_anexos,
            i_chunk,
            total,
            texto_hints=hints,
            kb_context=kb_ctx
        )

    t1 = _t()
    with ThreadPoolExecutor(max_workers=max(1, ANALISIS_CONCURRENCY)) as ex:
        futs = [ex.submit(_work, i+1, len(partes), partes[i]) for i in range(len(partes))]
        for fut in as_completed(futs):
            try:
                parciales.append(fut.result())
            except Exception as e:
                parciales.append(f"[ERROR parcial] {e}")

    _log_tiempo("parciales_multi_pass", t1)

    # Consolidación
    t2 = _t()
    borrador = _agregar_y_consolidar(parciales, varios_anexos, texto_hints=hints, kb_context=kb_ctx)
    final = _postproceso_final(borrador, raw, varios_anexos)
    _log_tiempo("consolidacion_multi_pass", t2)
    _log_tiempo("pipeline_multi_pass_total", t0)
    return final


# ==================== Exportar a PDF (ReportLab) ====================
def _wrap_lines(s: str, max_chars: int = 110) -> List[str]:
    """
    # Wrap simple por caracteres (word-wrap), suficiente para A4 con márgenes y 10pt.
    """
    out: List[str] = []
    for ln in (s or "").splitlines():
        w = ln.strip("\r")
        if not w:
            out.append("")
            continue
        parts: List[str] = []
        buf: List[str] = []
        for tok in re.findall(r"\S+|\s+", w):
            if tok.isspace():
                if sum(len(x) for x in buf) + len(tok) > max_chars:
                    parts.append("".join(buf).rstrip())
                    buf = []
                else:
                    buf.append(tok)
            else:
                if sum(len(x) for x in buf) + len(tok) > max_chars:
                    if buf:
                        parts.append("".join(buf).rstrip())
                        buf = [tok]
                    else:
                        parts.append(tok[:max_chars])
                        buf = [tok[max_chars:]]
                else:
                    buf.append(tok)
        if buf:
            parts.append("".join(buf).rstrip())
        out.extend(parts if parts else [""])
    return out

def generar_pdf_informe(texto_markdown: str, out_path: Optional[str] = None) -> str:
    """
    Crea un PDF simple y legible desde el texto del informe.
    - Convierte markdown light -> texto plano con bullets.
    - Usa márgenes y salto de página autom.
    Devuelve la ruta del PDF generado.
    """
    # Normaliza a texto plano para PDF
    contenido = preparar_texto_para_pdf(texto_markdown o "")

    # Si no pasaron ruta, generamos una con el formato que espera /descargar:
    # resumen_YYYYMMDDHHMMSS.pdf  (SIN guión bajo entre fecha y hora)
    if not out_path:
        try:
            tz = ZoneInfo("America/Argentina/Buenos_Aires")
        except Exception:
            tz = None
        ts = datetime.now(tz=tz).strftime("%Y%m%d%H%M%S")
        out_path = os.path.abspath(f"resumen_{ts}.pdf")

    # (opcional) log visible en Render para debug
    print(f"[PDF] generado en: {out_path}")

    c = canvas.Canvas(out_path, pagesize=A4)
    width, height = A4
    left = 20 * mm
    right = 15 * mm
    top = 18 * mm
    bottom = 18 * mm

    usable_w = width - left - right
    x = left
    y = height - top

    c.setTitle("Informe generado")
    c.setAuthor("Pliegos AI")
    c.setSubject("Informe técnico-jurídico")

    c.setFont("Helvetica-Bold", 12)
    c.setFillColor(HexColor("#222222"))
    c.drawString(x, y, "Informe técnico-jurídico")
    y -= 8 * mm

    c.setFont("Helvetica", 10)
    c.setFillColor(HexColor("#000000"))

    # Línea separadora
    c.line(x, y, x + usable_w, y)
    y -= 6 * mm

    # Cuerpo
    lines = _wrap_lines(contenido, max_chars=110)
    for ln in lines:
        if y < bottom + 15 * mm:
            c.showPage()
            y = height - top
            c.setFont("Helvetica", 10)
            c.setFillColor(HexColor("#000000"))
        c.drawString(x, y, ln)
        y -= 5 * mm

    c.showPage()
    c.save()
    return out_path

# ==================== Compat extra (plantillas) ====================
def generar_pdf_con_plantilla(
    informe_texto: str,
    *,
    plantilla: Optional[str] = None,  # se ignora, mantenida por compatibilidad
    salida: Optional[str] = None,
    **kwargs,
) -> str:
    """
    Compatibilidad con versiones anteriores que pedían generar PDF usando 'plantillas'.
    Ignora 'plantilla' y delega al generador simple.
    """
    return generar_pdf_informe(informe_texto, out_path=salida)


# ==================== Helpers de alto nivel ====================
def generar_informe_y_pdf(
    texto_fuente: str,
    *,
    varios_anexos: Optional[bool] = None,
    force_multi: Optional[bool] = None,
    export_pdf: bool = True,
    ruta_pdf: Optional[str] = None,
    prefer_source: Optional[str] = None,
) -> Tuple[str, Optional[str]]:
    """
    Atajo: corre el pipeline de análisis y, opcionalmente, exporta a PDF.
    Devuelve (informe_texto, ruta_pdf | None)
    """
    informe = analizar_y_generar_informe(
        texto_fuente,
        varios_anexos=varios_anexos,
        force_multi=force_multi,
        prefer_source=prefer_source,
    )
    # Si el input estaba incompleto, no generes PDF
    if informe.strip().startswith("ERROR: texto de entrada incompleto"):
        return informe, None

    pdf_path = generar_pdf_informe(informe, out_path=ruta_pdf) if export_pdf else None
    return informe, pdf_path

# ==================== Chat general (asistente) con KB ====================
def responder_chat_openai(
    prompt_or_messages,
    *,
    model: Optional[str] = None,
    max_tokens: Optional[int] = None,
    temperature: Optional[float] = None,
    system: Optional[str] = None,
    tools: Optional[List[Dict[str, Any]]] = None,
    tool_choice: Optional[Any] = None,
    prefer_source: Optional[str] = None,
    **kwargs,
) -> str:
    """
    Chat general orientado a licitaciones, con acceso a KB.
    - Si hay KB, se inyecta como CONTEXTO NO-CITABLE (guía, no fuente).
    - Respeta el prompt de chat en prompts.py (CHAT_ASSISTANT_PROMPT).
    """
    mod = _get_prom()
    chat_sys = ""
    if mod and hasattr(mod, "CHAT_ASSISTANT_PROMPT"):
        chat_sys = getattr(mod, "CHAT_ASSISTANT_PROMPT") or ""
    base_system = (system or chat_sys or "")

    # Normaliza mensajes
    if isinstance(prompt_or_messages, str):
        query_for_kb = prompt_or_messages[:1000]
        kb_ctx = build_kb_context(query=query_for_kb, source=prefer_source, top_k=8)
        messages = []
        sys = base_system
        if kb_ctx:
            sys = (sys + "\n\n" + "=== CONTEXTO KB (no-citable) ===\n" +
                   "Usar solo como guía; no inventar ni citar datos inexistentes.\n" + kb_ctx).strip()
        if sys:
            messages.append({"role": "system", "content": sys})
        messages.append({"role": "user", "content": prompt_or_messages})
    else:
        msgs = list(prompt_or_messages or [])
        last_user = next((m.get("content", "") for m in reversed(msgs) if m.get("role") == "user"), "")
        kb_ctx = build_kb_context(query=(last_user or "")[:1000], source=prefer_source, top_k=8)
        messages = msgs[:]
        if kb_ctx or base_system:
            sys_block = base_system
            if kb_ctx:
                sys_block = (sys_block + "\n\n" + "=== CONTEXTO KB (no-citable) ===\n" +
                             "Usar solo como guía; no inventar ni citar datos inexistentes.\n" + kb_ctx).strip()
            if not messages or messages[0].get("role") != "system":
                messages = [{"role": "system", "content": sys_block}] + messages
            else:
                messages[0]["content"] = (messages[0].get("content") or "") + "\n\n" + sys_block

    # Modelo y temperatura
    use_model = (model or _pick_model(final_pass=False))
    temp = temperature if (temperature is not None) else _mk_temperature()

    payload: Dict[str, Any] = {"model": use_model, "messages": messages}
    if temp is not None:
        payload["temperature"] = float(temp)
    if max_tokens is not None:
        payload["max_completion_tokens"] = int(max_tokens)
    if tools:
        payload["tools"] = tools
    if tool_choice:
        payload["tool_choice"] = tool_choice

    # _chat_create_safe devuelve str directamente
    out = _chat_create_safe(**payload)
    return (out or "").strip()

# ==================== Compat/alias ====================
def analizar_con_openai(
    texto_fuente: str,
    *,
    varios_anexos: Optional[bool] = None,
    force_multi: Optional[bool] = None,
    prefer_source: Optional[str] = None,
    **kwargs,
) -> str:
    """Alias legacy hacia analizar_y_generar_informe (con KB opcional)."""
    return analizar_y_generar_informe(
        texto_fuente,
        varios_anexos=varios_anexos,
        force_multi=force_multi,
        prefer_source=prefer_source,
    )

def generar_pdf(informe_texto: str, ruta_pdf: Optional[str] = None) -> str:
    """Alias histórico para exportar a PDF."""
    return generar_pdf_informe(informe_texto, out_path=ruta_pdf)

def analizar_y_pdf(
    texto_fuente: str,
    *,
    varios_anexos: Optional[bool] = None,
    force_multi: Optional[bool] = None,
    ruta_pdf: Optional[str] = None,
    prefer_source: Optional[str] = None,
) -> Tuple[str, Optional[str]]:
    """Alias cómodo equivalente a generar_informe_y_pdf."""
    return generar_informe_y_pdf(
        texto_fuente,
        varios_anexos=varios_anexos,
        force_multi=force_multi,
        export_pdf=True,
        ruta_pdf=ruta_pdf,
        prefer_source=prefer_source,
    )

# ==================== __all__ ====================
try:
    __all__
except NameError:
    __all__ = []

try:
    __all__.extend([
        # nuevas
        "analizar_y_generar_informe", "generar_informe_y_pdf", "generar_pdf_informe",
        # compat
        "analizar_con_openai", "analizar_y_pdf", "generar_pdf", "generar_pdf_con_plantilla",
        # chat
        "responder_chat_openai",
        # RAG helpers
        "build_kb_context",
    ])
except Exception:
    pass
